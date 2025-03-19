from io import BytesIO
import docx
from langchain_core.documents import Document
import logging
from pathlib import Path
import tempfile
from src.utils.document_loaders import IncomingFileProcessor
from src.utils.logger import get_logger

logger = get_logger(__name__)

async def handling_documents(contents, file_extension, original_filename):
    file_processor = IncomingFileProcessor(chunk_size=512)
    try:
        if contents:
            suffix = "docx" if file_extension.endswith('docx') else "pdf"

            logger.info("Creating temporary file for the uploaded file")
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(contents)

            try:
                if file_extension.endswith('docx'):
                    texts = file_processor.get_docx_splits(tmp_file.name, original_filename)
                    logger.info("Successfully got docx chunks")

                elif file_extension.endswith('pdf'):
                    texts = file_processor.get_pdf_splits(tmp_file.name, original_filename)
                    logger.info("Successfully got pdf chunks")

            except Exception as e:
                logger.critical(f"Error in loading/splitting uploaded file: {e}")
                print(e)

            finally:
                # Clean up the temporary file
                logger.info("Cleaning the temporary file")
                Path(tmp_file.name).unlink()

            return texts
        
    except Exception as e:
        logger.critical(f"Error while handling document: {e}")
        raise Exception(f"Error while handling document: {e}")

# async def extract_text(file, file_extension: str) -> str:
#     """Extract text based on file type."""
#     if file_extension == "pdf":
#         pdf_document = PyMuPDFLoader(file)
#         return "\n".join([pdf_document[page].get_text("text") for page in range(len(pdf_document))])
#     elif file_extension in ["doc", "docx"]:
#         doc = docx.Document(BytesIO(file))
#         return "\n".join([para.text for para in doc.paragraphs])
#     elif file_extension == "txt":
#         return file.decode("utf-8")
#     else:
#         raise ValueError("Unsupported file format.")

# def split_text(text: str, filename: str):
#     """Splits text into chunks."""
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#     return [Document(page_content=chunk, metadata={"source": filename}) for chunk in text_splitter.split_text(text)]
